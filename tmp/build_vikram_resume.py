from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"E:\xampp\htdocs\himalayanwaterscience\output\docx\Vikram_Singh_Rathore_Senior_Backend_Resume.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "17365D"
BLUE = "1F4E79"
INK = "202B38"
MUTED = "586574"
LIGHT = "D9E2F3"


def set_font(run, size=9.3, bold=False, italic=False, color=INK):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url, color=MUTED):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    rpr.append(fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rpr.append(sz)
    new_run.append(rpr)
    txt = OxmlElement("w:t")
    txt.text = text
    new_run.append(txt)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.48)
sec.bottom_margin = Inches(0.48)
sec.left_margin = Inches(0.58)
sec.right_margin = Inches(0.58)
sec.header_distance = Inches(0.25)
sec.footer_distance = Inches(0.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(9.3)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(2.5)
normal.paragraph_format.line_spacing = 1.03

bullet = styles["List Bullet"]
bullet.font.name = "Arial"
bullet._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
bullet.font.size = Pt(9.1)
bullet.paragraph_format.left_indent = Inches(0.22)
bullet.paragraph_format.first_line_indent = Inches(-0.14)
bullet.paragraph_format.space_after = Pt(1.5)
bullet.paragraph_format.line_spacing = 1.02

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(1)
set_font(title.add_run("VIKRAM SINGH RATHORE"), 20.5, True, color=NAVY)

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.paragraph_format.space_after = Pt(3)
set_font(tag.add_run("SENIOR BACKEND ENGINEER | TECH LEAD | NODE.JS & PHP/LARAVEL"), 10.3, True, color=BLUE)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(6)
set_font(contact.add_run("+91-7014463427  |  vikramrathore66223@gmail.com  |  "), 9, color=MUTED)
add_hyperlink(contact, "linkedin.com/in/vikram-singh-rathore-999bb7134", "https://www.linkedin.com/in/vikram-singh-rathore-999bb7134/")
set_font(contact.add_run("  |  Udaipur, India"), 9, color=MUTED)


def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_font(r, 10.5, True, color=BLUE)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), LIGHT)
    pBdr.append(bottom)
    pPr.append(pBdr)


def body(text, after=2.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.04
    set_font(p.add_run(text), 9.3)
    return p


def skill_line(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.2)
    set_font(p.add_run(label + ": "), 9.1, True, color=NAVY)
    set_font(p.add_run(value), 9.1)


def role(title_text, dates, company, location):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.25), WD_TAB_ALIGNMENT.RIGHT)
    set_font(p.add_run(title_text), 10, True, color=NAVY)
    set_font(p.add_run("\t" + dates), 9.4, True, color=MUTED)
    p2 = doc.add_paragraph()
    p2.paragraph_format.keep_with_next = True
    p2.paragraph_format.space_after = Pt(1.5)
    p2.paragraph_format.tab_stops.add_tab_stop(Inches(7.25), WD_TAB_ALIGNMENT.RIGHT)
    set_font(p2.add_run(company), 9.3, True)
    set_font(p2.add_run("\t" + location), 9.1, italic=True, color=MUTED)


def item(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    set_font(p.add_run(text), 9.05)


section_heading("Professional Summary")
body("Senior Backend Engineer and Tech Lead with 8+ years of experience delivering production systems across Node.js/Express.js and PHP/Laravel. Strong background in REST API design, backend architecture, MySQL/RDS optimization, authentication and authorization, payment workflows, AWS deployments, ERP/e-commerce platforms, and third-party integrations. Owns delivery from requirements and technical design through implementation, production release, troubleshooting, and team mentoring.")

section_heading("Core Technical Skills")
skill_line("Backend", "Node.js, Express.js, JavaScript, PHP, Laravel, REST APIs, Webhooks, Microservices")
skill_line("Data", "MySQL, Amazon RDS, Prisma ORM; working exposure to MongoDB")
skill_line("Cloud & Production", "AWS EC2, S3 and RDS, Linux, PM2, Apache, Git")
skill_line("Security & Payments", "OTP authentication, JWT, RBAC, Razorpay, GST workflows")
skill_line("Integrations", "SAP ERP, Zendesk, WhatsApp/SMS, POS, E-Invoice, E-Waybill, Gemini and OpenAI APIs")
skill_line("Additional", "Next.js, React, Python, FastAPI, SQL, Jira")

section_heading("Professional Experience")
role("Tech Lead / Team Lead - Backend & Full Stack", "Feb 2025 - Present", "Gyanitalk Technologies (OPC) Private Limited", "Udaipur, India")
item("Lead end-to-end engineering of an astrology and consultation platform, covering requirement analysis, backend architecture, API design, implementation, production deployment, and cross-functional delivery.")
item("Build and maintain backend services using Node.js, Express.js, PHP/Laravel, Prisma ORM, and MySQL within a mixed-service architecture.")
item("Designed secure authentication and authorization flows using OTP login, JWT tokens, and role-based access control for customer and operational workflows.")
item("Implemented Razorpay payment flows, webhook handling, GST calculation, and automated invoice generation for consultation transactions.")
item("Built a Python/FastAPI service that processes structured Kundli data and prepares context for Gemini and OpenAI-powered AI Astrologer workflows.")
item("Improve API and database performance through query review and Prisma/MySQL optimization; mentor developers and coordinate delivery with product, frontend, and QA teams.")

role("Senior Backend Developer", "Apr 2022 - Jan 2025", "Wooden Street Furnitures Pvt. Ltd.", "Udaipur, India")
item("Developed and maintained backend systems for high-volume e-commerce and ERP workflows using Laravel, Node.js, MySQL/RDS, REST APIs, and AWS infrastructure.")
item("Delivered APIs for product catalog, cart, checkout, order management, manufacturing, inventory, accounts, reporting, and order tracking.")
item("Integrated SAP ERP, Zendesk, payment gateways, WhatsApp/SMS services, POS systems, GST E-Invoice, and E-Waybill workflows.")
item("Improved database performance and production stability through SQL query optimization, indexing, and backend troubleshooting.")
item("Worked across business-critical modules and integration boundaries to translate operational requirements into reliable backend workflows.")

role("Senior Backend Developer", "Aug 2018 - Mar 2022", "Cognus Technology", "Udaipur, India")
item("Developed backend APIs and authentication systems for multiple client applications using Node.js, PHP/Laravel, and MySQL.")
item("Implemented integrations with payment gateways, SMS providers, and analytics services, including API contracts and error-handling workflows.")
item("Optimized database queries and API endpoints and collaborated with frontend engineers on clear, maintainable API contracts.")

section_heading("Selected Engineering Work")
role("Astrology, Consultation & AI Platform", "2025", "Node.js, Express.js, Laravel, Prisma, MySQL, FastAPI", "")
item("Delivered consultation, authentication, payments, GST invoicing, and AI-assisted astrology workflows across Node.js, Laravel, and FastAPI services.")
item("Integrated Gemini and OpenAI services through structured context preparation while keeping core platform workflows in the broader backend ecosystem.")

role("E-Commerce & Multi-Module ERP", "2022 - 2025", "Laravel, Node.js, MySQL/RDS, AWS, REST APIs", "")
item("Built and enhanced connected commerce and ERP workflows spanning catalog, checkout, orders, manufacturing, inventory, accounts, reporting, and external systems.")

section_heading("Education")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(1)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.tab_stops.add_tab_stop(Inches(7.25), WD_TAB_ALIGNMENT.RIGHT)
set_font(p.add_run("Bachelor of Technology (B.Tech), Computer Science Engineering"), 9.4, True, color=NAVY)
set_font(p.add_run("\t2013 - 2017"), 9.1, True, color=MUTED)
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(0)
set_font(p2.add_run("Rajasthan Technical University - SS College of Engineering, Udaipur"), 9.1)

doc.core_properties.title = "Vikram Singh Rathore - Senior Backend Engineer Resume"
doc.core_properties.subject = "Senior Node.js and PHP/Laravel Backend Engineer"
doc.core_properties.author = "Vikram Singh Rathore"
doc.core_properties.keywords = "Node.js, Express.js, PHP, Laravel, Backend Engineer, Tech Lead, AWS, MySQL"
doc.save(OUT)
print(OUT)
